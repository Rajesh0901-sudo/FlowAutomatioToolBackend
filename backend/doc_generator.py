from docx import Document
from check_existing_file import check_file_status

def save_results_to_docx(column_names, rows, file_name="query_results.docx"):

    if file_name == (None, ''):
        response = 'File name is empty'
        print(response)
        return 0
    
    elif column_names == (None, ''):
        response = 'Header name is empty'
        print(response)
        return 0
    
    elif rows == (None, ''):
        response = 'No Data Found to Save in Doc'
        print(response)
        return 0

    if_exisiting_doc = check_file_status(file_name)

    if if_exisiting_doc == 1:
        doc = Document(file_name)
    elif if_exisiting_doc == 2:
        doc = Document()
    else:
        print("Got and error")
        return 0


    doc.add_heading("Query Results", level=1)

    table = doc.add_table(rows=1,cols=len(column_names))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells

    for idx,column_name in enumerate(column_names):
        hdr_cells[idx].text = column_name

    for row in rows:
        row_cells = table.add_row().cells
        for idx,data in enumerate(row):
            row_cells[idx].text = str(data)

    doc.save(file_name)
    print("Document saved")
    return 1