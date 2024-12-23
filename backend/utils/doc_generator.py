from docx import Document
from docx.shared import Pt,Mm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from utils.check_existing_file import check_file_status


def save_results_to_docx(db_name,tableName,column_names, rows, file_name="query_results.docx"):

    if file_name == (None, ''):
        response = 'File name is empty'
        print(response)
        return 0
    
    elif column_names == (None, ''):
        response = 'Header name is empty'
        print(response)
        return 0
    
    elif rows == (None, ''):
        response = 'No Data Found to Save in Doc'
        print(response)
        return 0
    

    if_exisiting_doc = check_file_status(file_name)


    if if_exisiting_doc == 1:
        doc = Document(file_name)
    elif if_exisiting_doc == 2:
        doc = Document()
    else:
        print("Got and error")
        return 0
    
    styles = doc.styles
    styles['Heading 3'].font.color.rgb = RGBColor(0, 0, 0)
    
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(15.4)
    section.right_margin = Mm(15.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    heading = doc.add_heading( tableName.upper() + ' Query:-', level=5)
    heading.paragraph_format.space_before = Mm(8.4)
    heading.paragraph_format.space_after = Mm(2.4)

    table = doc.add_table(rows=1,cols=len(column_names))
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells

    for idx,column_name in enumerate(column_names):
        hdr_cells[idx].text = column_name.upper()

        run = hdr_cells[idx].paragraphs[0].runs[0]
        run.font.size = Pt(8)
        #set_cell_background_color(hdr_cells[idx],'0000FF')

    for row in rows:
        row_cells = table.add_row().cells
        for idx,data in enumerate(row):
            row_cells[idx].text = str(data)
            run = row_cells[idx].paragraphs[0].runs[0]
            run.font.size = Pt(8)

    doc.save(file_name)
    print("Document saved")
    return 1

def set_cell_background_color(cell,color):
    tc = cell.element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'),color)
    tcPr.append(shd)